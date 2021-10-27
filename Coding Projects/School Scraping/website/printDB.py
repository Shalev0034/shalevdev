from .models import User,Table,Major
from . import db
import docx
import os
from docx.shared import Pt
import pyperclip
import sys


def example(name,query,cl,doc):
    # Create an instance of a word document
  
    # Add a Title to the document
    doc.add_heading(name, 0)    
    
    #Get the vars of the table
    lst = getVars(cl)

    # Creating a table object
    table = doc.add_table(rows=1, cols=len(lst))
    table.autofit=False

    # Adding heading in the 1st row of the table

    row = table.rows[0].cells

    for i in range(len(lst)):
        row[i].text=[lst[i]]

     
    # Adding data from the list to the table
    for data in query:
  
        # Adding a row and then adding data in it.
        row = table.add_row().cells

        #vars(user) returns an unordered list of the values, this function orders them
        values = orderValues(lst,vars(data))

      
        #dinamically change to font size according ot the size of the value
        for i in range(len(values)):
            
            par = row[i].add_paragraph(str(values[i]))

            length = len(str(values[i]))

            if len(par.runs)>0:
                font = par.runs[0].font

                if length>300:
                    font.size= Pt(1)
                elif length>86:
                    font.size= Pt(5)
                elif length>24:
                    font.size = Pt(8)
                else:
                    font.size = Pt(11)     
  
  
    # Now save the document to a location
    os.remove(sys.path[0]+"/f.docx")

    doc.save(sys.path[0]+"/f.docx")

def printDB(): 
    doc = docx.Document()

    example("Users",User.query.all(),User,doc)  
    space(doc)
    example("Major",Major.query.all(),Major,doc)
    space(doc)
    example("Tables",Table.query.all(),Table,doc)

    
def space(doc):
    for i in range(3):
        doc.add_paragraph()

def getVars(cl):
    lst = []
    for var in vars(cl):
        if var != '__doc__':
            lst.append(var)
        else: break
    lst.remove('__module__')
    return lst
    

def orderValues(lst,items):
    newValues=[]
    for key in lst:
        newValues.append(items.get(key))
    return newValues
